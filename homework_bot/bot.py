import asyncio
import os
import io
import re
from dotenv import load_dotenv

from aiogram import Bot, Dispatcher, types, F
from aiogram.client.session.aiohttp import AiohttpSession
from aiogram.filters import CommandStart, Command
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton, BufferedInputFile

# مكتبات تنسيق مستندات Word
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# مكتبات تنسيق مستندات PDF
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# مكتبات معالجة تشبيك واتجاه النصوص العربية (BiDi)
try:
    import arabic_reshaper
    from bidi.algorithm import get_display
    ARABIC_LIB_INSTALLED = True
except ImportError:
    ARABIC_LIB_INSTALLED = False

import database as db
import ai_solver

load_dotenv()

BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
ADMIN_ID = int(os.getenv("ADMIN_ID", "0"))

session = AiohttpSession(timeout=60)
bot = Bot(token=BOT_TOKEN, session=session)
dp = Dispatcher()

ACTIVE_SUPPORT = set()
USER_LAST_TEXT = {}
USER_PENDING_TASKS = {}

# محاولة تسجيل خط عربي يدعم تقارير الـ PDF لمنع ظهور المربعات السوداء
ARABIC_FONT_NAME = "Helvetica"
for font_path in ["Cairo-Regular.ttf", "Amiri-Regular.ttf", "Arial.ttf", "fonts/Cairo-Regular.ttf"]:
    if os.path.exists(font_path):
        try:
            pdfmetrics.registerFont(TTFont("ArabicFont", font_path))
            ARABIC_FONT_NAME = "ArabicFont"
            break
        except Exception:
            pass


# --- دوال توحيد النصوص والاتجاه (RTL / LTR) ---
def contains_arabic(text: str) -> bool:
    """التحقق من احتواء النص على حروف عربية"""
    return bool(re.search(r'[\u0600-\u06FF]', text))


def normalize_arabic(text: str) -> str:
    """توحيد أشكال الهمزات والتاء المربوطة والألف المقصورة"""
    if not text:
        return ""
    text = text.lower()
    text = re.sub(r"[إأآٱ]", "ا", text)
    text = re.sub(r"ة", "ه", text)
    text = re.sub(r"ى", "ي", text)
    text = re.sub(r"[\u064B-\u065F]", "", text)
    return text


def set_paragraph_direction(paragraph, is_arabic: bool):
    """ضبط اتجاه الفقرة في Word من اليمين لليسار عند وجود عربي"""
    if is_arabic:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pPr = paragraph._p.get_or_add_pPr()
        pPr.append(parse_xml(r'<w:bidi %s/>' % nsdecls('w')))
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def format_bidi_text(text: str) -> str:
    """إعادة تشكيل وترتيب الكلمات العربية المتداخلة مع الإنجليزية للـ PDF"""
    if not text or not ARABIC_LIB_INSTALLED or not contains_arabic(text):
        return text
    try:
        configuration = {
            'delete_harakat': False,
            'support_ligatures': True,
        }
        reshaper = arabic_reshaper.ArabicReshaper(configuration=configuration)
        reshaped_text = reshaper.reshape(text)
        return get_display(reshaped_text)
    except Exception:
        return text


SUPPORT_KEYWORDS = [
    "دعم", "خدمه العملاء", "مشكل", "مساعد", "استفسار",
    "اداره", "مشرف", "شخص", "اتكلم", "تواصل", "خدمه",
    "support", "help", "admin"
]


# --- استخراج وقراءة الملفات المرفوعة ---
def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        extracted = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_data = [c.text.strip() for c in row.cells if c.text.strip()]
                if row_data:
                    extracted.append(" | ".join(row_data))
        return "\n".join(extracted)
    except Exception:
        return ""


# --- تحويل جداول الماركداون إلى جداول Word حقيقية ومنسقة ---
def add_markdown_table_to_docx(doc: Document, table_lines: list):
    rows_data = []
    for line in table_lines:
        stripped = line.strip()
        # تجاهل خطوط الفصل مثل |---|---|
        if re.match(r"^\|?\s*[-:]+\s*(\|\s*[-:]+\s*)+\|?$", stripped):
            continue
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        if any(cells):
            rows_data.append(cells)

    if not rows_data:
        return

    cols_count = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=cols_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row in enumerate(rows_data):
        is_header = (r_idx == 0)
        for c_idx in range(cols_count):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text
            
            tc_pr = cell._tc.get_or_add_tcPr()
            if is_header:
                # ترويسة الجدول بلون كحلي أكاديمي وكتابة بيضاء
                shd = parse_xml(r'<w:shd %s w:fill="1B365D"/>' % nsdecls("w"))
                tc_pr.append(shd)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(10)
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                for p in cell.paragraphs:
                    if contains_arabic(cell_text):
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in p.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(9.5)

    doc.add_paragraph()


# --- توليد ملف Word أكاديمي نظيف 100% ---
def create_pro_academic_docx(solution_text: str, custom_title: str, filename: str) -> BufferedInputFile:
    doc = Document()
    
    # ضبط الهوامش وإضافة أرقام الصفحات (بدون أي ترويسة تدل على البوت)
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

        # إضافة ترقيم الصفحات في الأسفل
        footer = s.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_p.add_run()
        footer_run.font.name = "Arial"
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        fld = parse_xml(r'<w:fldSimple %s w:instr="PAGE"/>' % nsdecls("w"))
        footer_p._p.append(fld)

    # عنوان المستند الرئيسي
    is_title_ar = contains_arabic(custom_title)
    title_p = doc.add_paragraph()
    set_paragraph_direction(title_p, is_title_ar)
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(18)
    
    title_run = title_p.add_run(custom_title)
    title_run.font.name = "Arial"
    title_run.font.size = Pt(16)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(27, 54, 93)
    if is_title_ar:
        title_run._r.get_or_add_rPr().append(parse_xml(r'<w:rtl %s/>' % nsdecls("w")))

    lines = solution_text.split("\n")
    table_buffer = []

    for line in lines:
        raw = line.strip()

        # كشف الجداول وتجميع أسطرها
        if raw.startswith("|") and raw.endswith("|"):
            table_buffer.append(raw)
            continue
        else:
            if table_buffer:
                add_markdown_table_to_docx(doc, table_buffer)
                table_buffer = []

        if not raw:
            continue

        is_ar = contains_arabic(raw)
        cleaned = raw.replace("$$", "").replace("$", "").replace("\\text{", "").replace("}", "")

        # العناوين
        if raw.startswith(("### ", "## ", "# ")):
            hp = doc.add_paragraph()
            set_paragraph_direction(hp, is_ar)
            hp.paragraph_format.space_before = Pt(12)
            hp.paragraph_format.space_after = Pt(4)
            hrun = hp.add_run(raw.lstrip("#").strip().replace("**", ""))
            hrun.font.name = "Arial"
            hrun.font.size = Pt(12.5)
            hrun.bold = True
            hrun.font.color.rgb = RGBColor(41, 70, 115)
            if is_ar:
                hrun._r.get_or_add_rPr().append(parse_xml(r'<w:rtl %s/>' % nsdecls("w")))

        # القوائم النقطية
        elif raw.startswith(("- ", "* ", "• ")):
            bp = doc.add_paragraph(style="List Bullet")
            set_paragraph_direction(bp, is_ar)
            bp.paragraph_format.space_after = Pt(3)
            item_text = cleaned.lstrip("-*• ").strip()
            parts = item_text.split("**")
            for i, part in enumerate(parts):
                brun = bp.add_run(part)
                brun.font.name = "Arial"
                brun.font.size = Pt(10.5)
                if is_ar:
                    brun._r.get_or_add_rPr().append(parse_xml(r'<w:rtl %s/>' % nsdecls("w")))
                if i % 2 == 1:
                    brun.bold = True
                    brun.font.color.rgb = RGBColor(27, 54, 93)

        # الفقرات العادية
        else:
            p = doc.add_paragraph()
            set_paragraph_direction(p, is_ar)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            parts = cleaned.split("**")
            for i, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = "Arial"
                run.font.size = Pt(10.5)
                if is_ar:
                    run._r.get_or_add_rPr().append(parse_xml(r'<w:rtl %s/>' % nsdecls("w")))
                if i % 2 == 1:
                    run.bold = True
                    run.font.color.rgb = RGBColor(27, 54, 93)

    if table_buffer:
        add_markdown_table_to_docx(doc, table_buffer)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return BufferedInputFile(file_stream.getvalue(), filename=f"{filename}.docx")


# --- توليد مستند PDF أكاديمي ---
def create_pro_academic_pdf(solution_text: str, custom_title: str, filename: str) -> BufferedInputFile:
    pdf_stream = io.BytesIO()
    doc = SimpleDocTemplate(pdf_stream, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    
    font_main = ARABIC_FONT_NAME
    
    title_style = ParagraphStyle(
        "DocTitle",
        parent=styles["Heading1"],
        fontName=font_main,
        fontSize=15,
        leading=20,
        textColor=colors.HexColor("#1B365D"),
        alignment=1,
        spaceAfter=14
    )
    
    story = [Paragraph(format_bidi_text(custom_title), title_style), Spacer(1, 6)]
    
    for line in solution_text.split("\n"):
        raw = line.strip()
        if not raw:
            continue
            
        is_ar = contains_arabic(raw)
        align_choice = 2 if is_ar else 0  # 2 = محاذاة يمين للعربي، 0 = محاذاة يسار للإنجليزي
        
        heading_style = ParagraphStyle(
            "HeadingStyle",
            parent=styles["Heading2"],
            fontName=font_main,
            fontSize=12,
            leading=16,
            textColor=colors.HexColor("#294673"),
            alignment=align_choice,
            spaceBefore=10,
            spaceAfter=4
        )
        
        body_style = ParagraphStyle(
            "BodyStyle",
            parent=styles["Normal"],
            fontName=font_main,
            fontSize=10,
            leading=15,
            textColor=colors.HexColor("#252525"),
            alignment=align_choice,
            spaceAfter=5
        )

        cleaned = raw.replace("$$", "").replace("$", "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        if raw.startswith(("### ", "## ", "# ")):
            h_text = raw.lstrip("#").strip().replace("**", "")
            story.append(Paragraph(f"<b>{format_bidi_text(h_text)}</b>", heading_style))
        else:
            parts = cleaned.split("**")
            formatted_parts = []
            for i, p in enumerate(parts):
                bidi_p = format_bidi_text(p)
                formatted_parts.append(f"<b>{bidi_p}</b>" if i % 2 == 1 else bidi_p)
            story.append(Paragraph("".join(formatted_parts), body_style))
        
    doc.build(story)
    pdf_stream.seek(0)
    return BufferedInputFile(pdf_stream.getvalue(), filename=f"{filename}.pdf")


# --- لوحات الأزرار ---
def get_main_menu_keyboard():
    return InlineKeyboardMarkup(inline_keyboard=[
        [
            InlineKeyboardButton(text="💬 خدمة العملاء والدعم", callback_data="btn_support"),
        ]
    ])


def get_options_keyboard():
    return InlineKeyboardMarkup(inline_keyboard=[
        [
            InlineKeyboardButton(text="🇬🇧 Word (.docx)", callback_data="opt_en_docx"),
            InlineKeyboardButton(text="🇬🇧 PDF (.pdf)", callback_data="opt_en_pdf"),
        ],
        [
            InlineKeyboardButton(text="🇸🇦 وورد (.docx)", callback_data="opt_ar_docx"),
            InlineKeyboardButton(text="🇸🇦 بي دي إف (.pdf)", callback_data="opt_ar_pdf"),
        ],
        [
            InlineKeyboardButton(text="💬 نص مباشر", callback_data="opt_en_txt"),
        ]
    ])


# --- معالجة الأوامر ---
@dp.message(CommandStart())
async def cmd_start(message: types.Message):
    user_id = message.from_user.id
    ACTIVE_SUPPORT.discard(user_id)
    credits = await db.get_or_create_user(user_id)
    
    welcome_text = (
        f"يا هلا والله بـ **{message.from_user.first_name}**، نوّرت! 🎓✨\n\n"
        "أنا مساعدك الجامعي الذكي، موجود هنا عشان أسهّل عليك مشوارك وأساعدك في إنجاز الواجبات، تقارير اللاب، وتكاليفك الجامعية أول بأول.\n\n"
        "ولا تشيل هم التنسيق، حلك يوصلك مرتب وجاهز بملف **Word** أو **PDF** للتسليم فوراً! 😉\n\n"
        "────────────────\n"
        "📌 **كيف تبدأ بـ 3 خطوات بس؟**\n"
        "1️⃣ **اكتب شروطك** (اختياري) مثل: *إنجليزي فقط، ركز على الخاتمة والمراجع*.\n"
        "2️⃣ **أرسل ملف الواجب** على طول (صورة، PDF، أو Word).\n"
        "3️⃣ **اختر الصيغة** (**Word** أو **PDF**) واستلم ملفك فوراً!\n"
        "────────────────\n\n"
        f"🔹 **رصيدك الحالي:** {credits} نقاط.\n"
        "🎁 *(شحنا لك رصيد تجريبي مجاني عشان تجرب أول حل وتشوف الجودة بنفسك)*"
    )
    
    await message.answer(
        welcome_text,
        reply_markup=get_main_menu_keyboard(),
        parse_mode="Markdown"
    )


@dp.message(Command("add"))
async def cmd_add(message: types.Message):
    await db.add_credits(message.from_user.id, 10)
    await message.answer("🎁 تم شحن 10 نقاط في حسابك للتجربة!")


@dp.message(Command("give"))
async def cmd_give_credits(message: types.Message):
    user_id = message.from_user.id
    if user_id != ADMIN_ID:
        return

    args = message.text.split()
    if len(args) != 3:
        await message.answer(
            "⚠️ **طريقة الاستخدام:**\n"
            "`/give <Telegram_ID> <النقاط>`\n\n"
            "📌 **مثال:**\n"
            "`/give 769764499 50`",
            parse_mode="Markdown"
        )
        return

    try:
        target_id = int(args[1])
        amount = int(args[2])
        await db.add_credits(target_id, amount)
        
        await message.answer(
            f"✅ **تم بنجاح!**\nتمت إضافة **{amount}** نقطة للمستخدم: `{target_id}`.",
            parse_mode="Markdown"
        )
        try:
            await bot.send_message(
                chat_id=target_id,
                text=f"🎁 **مبروك!** تمت إضافة **{amount}** نقطة رصيد إلى حسابك في البوت 🎓✨\nيمكنك الآن حل تكاليفك بكل حرية!"
            )
        except Exception:
            pass

    except ValueError:
        await message.answer("❌ خطأ: يرجى التأكد من كتابة الآيدي وعدد النقاط كأرقام صحيحة.")


# --- إدارة خدمة العملاء ---
@dp.callback_query(F.data == "btn_support")
@dp.message(Command("support"))
async def trigger_support(event: types.CallbackQuery | types.Message):
    user = event.from_user
    ACTIVE_SUPPORT.add(user.id)
    
    msg_text = (
        "💬 **أنت الآن في محادثة مباشرة مع خدمة العملاء.**\n\n"
        "اكتب مشكلتك أو استفسارك وسيقوم المشرف بالرد عليك هنا مباشرة.\n\n"
        "🔙 للخروج من المحادثة والعودة لحل الواجبات اكتب: `/exit`"
    )
    
    if isinstance(event, types.CallbackQuery):
        await event.answer()
        await event.message.answer(msg_text, parse_mode="Markdown")
    else:
        await event.answer(msg_text, parse_mode="Markdown")

    if ADMIN_ID:
        try:
            await bot.send_message(
                chat_id=ADMIN_ID,
                text=f"🔔 **طلب دعم فني جديد!**\n\nالطالب: {user.full_name}\nالمعرف: `ID:{user.id}`\n\n(للرد عليه اضغط Reply على رسائله)",
                parse_mode="Markdown"
            )
        except Exception as e:
            print(f"Admin Notify Error: {e}")


@dp.message(Command("exit"))
async def exit_support(message: types.Message):
    user_id = message.from_user.id
    if user_id in ACTIVE_SUPPORT:
        ACTIVE_SUPPORT.remove(user_id)
        await message.answer("✅ تم إنهاء محادثة الدعم الفني، وعاد البوت للوضع التلقائي لحل الواجبات.")
        if ADMIN_ID:
            await bot.send_message(ADMIN_ID, f"🔒 تم إغلاق جلسة الدعم مع الطالب: `ID:{user_id}`")
    else:
        await message.answer("أنت لست في محادثة دعم حالياً.")


@dp.message(F.chat.id == ADMIN_ID, F.reply_to_message)
async def handle_admin_reply(message: types.Message):
    reply_text = message.reply_to_message.text or message.reply_to_message.caption or ""
    match = re.search(r"ID:(\d+)", reply_text)
    if match:
        target_user_id = int(match.group(1))
        try:
            if message.text:
                await bot.send_message(target_user_id, f"👨‍💼 **خدمة العملاء:**\n\n{message.text}", parse_mode="Markdown")
            elif message.photo:
                await bot.send_photo(target_user_id, photo=message.photo[-1].file_id, caption=f"👨‍💼 **خدمة العملاء:**\n\n{message.caption or ''}")
            elif message.document:
                await bot.send_document(target_user_id, document=message.document.file_id, caption=f"👨‍💼 **خدمة العملاء:**\n\n{message.caption or ''}")
            elif message.voice:
                await bot.send_voice(target_user_id, voice=message.voice.file_id, caption="👨‍💼 **تسجيل صوتي من الدعم الفني**")

            await message.reply("✅ تم إيصال ردك للطالب بنجاح.")
        except Exception as e:
            await message.reply(f"❌ تعذر إرسال الرد: {str(e)}")


# --- معالجة النصوص والملفات ---
@dp.message(F.text & ~F.text.startswith("/"))
async def handle_text(message: types.Message):
    user_id = message.from_user.id
    raw_text = message.text.strip()
    norm_text = normalize_arabic(raw_text)

    if user_id in ACTIVE_SUPPORT:
        if ADMIN_ID:
            await bot.send_message(
                ADMIN_ID,
                f"📩 **رسالة من الطالب:** {message.from_user.full_name} (`ID:{user_id}`)\n\n{raw_text}"
            )
            await message.answer("📨 تم إيصال رسالتك للدعم الفني، يرجى انتظار الرد...")
        return

    if any(k in norm_text for k in SUPPORT_KEYWORDS):
        await trigger_support(message)
        return

    USER_LAST_TEXT[user_id] = raw_text
    await message.answer("📝 **تم حفظ ملاحظاتك!**\nالآن أرسل ملف الواجب لتطبيقها على الحل فوراً.")


@dp.message(F.photo | F.document | F.voice)
async def handle_files(message: types.Message):
    user_id = message.from_user.id

    if user_id in ACTIVE_SUPPORT:
        if ADMIN_ID:
            caption = f"📎 **ملف من الطالب:** {message.from_user.full_name} (`ID:{user_id}`)\n{message.caption or ''}"
            if message.photo:
                await bot.send_photo(ADMIN_ID, photo=message.photo[-1].file_id, caption=caption)
            elif message.document:
                await bot.send_document(ADMIN_ID, document=message.document.file_id, caption=caption)
            elif message.voice:
                await bot.send_voice(ADMIN_ID, voice=message.voice.file_id, caption=caption)
            await message.answer("📨 تم إرسال المرفق للدعم الفني.")
        return

    credits = await db.get_or_create_user(user_id)
    if credits <= 0:
        await message.answer("⚠️ رصيدك 0 نقاط. تواصل مع الدعم أو اشحن رصيدك للمتابعة.")
        return

    file_bytes = None
    mime_type = "image/jpeg"
    caption = message.caption or ""
    prior_text = USER_LAST_TEXT.pop(user_id, "")
    text_query = f"{prior_text}\n{caption}".strip()

    if message.photo:
        file = await bot.get_file(message.photo[-1].file_id)
        f_io = io.BytesIO()
        await bot.download_file(file.file_path, destination=f_io)
        file_bytes = f_io.getvalue()
        mime_type = "image/jpeg"

    elif message.document:
        doc = message.document
        file = await bot.get_file(doc.file_id)
        f_io = io.BytesIO()
        await bot.download_file(file.file_path, destination=f_io)
        raw_bytes = f_io.getvalue()

        if (doc.file_name or "").lower().endswith(".docx"):
            docx_content = extract_text_from_docx(raw_bytes)
            text_query = f"{text_query}\n\n[Docx Content]:\n{docx_content}"
            file_bytes = None
        else:
            mime_type = doc.mime_type or "application/pdf"
            file_bytes = raw_bytes

    USER_PENDING_TASKS[user_id] = {
        "file_bytes": file_bytes,
        "mime_type": mime_type,
        "text_query": text_query
    }

    await message.answer("📥 **تم استلام الواجب!**\nحدد الصيغة المطلوبة لاستلام الحل:", reply_markup=get_options_keyboard())


@dp.callback_query(F.data.startswith("opt_"))
async def process_solution_options(callback: types.CallbackQuery):
    user_id = callback.from_user.id
    task = USER_PENDING_TASKS.get(user_id)

    if not task:
        await callback.answer("⚠️ أعد إرسال الملف من فضلك.", show_alert=True)
        return

    _, lang, out_type = callback.data.split("_")
    await callback.answer()
    
    wait_msg = await callback.message.edit_text("⏳ **جاري قراءة الملف وتحليل الحل بالذكاء الاصطناعي...**")

    # 1. إرسال الطلب لمحرك الذكاء الاصطناعي
    success, doc_title, safe_filename, solution = await ai_solver.solve_homework(
        text_query=task["text_query"],
        image_bytes=task["file_bytes"],
        mime_type=task["mime_type"],
        lang=lang
    )

    # إذا حدث خطأ (مثل 429 أو مشاكل كوتا): لا يتم الخصم
    if not success:
        await wait_msg.edit_text(solution)
        USER_PENDING_TASKS.pop(user_id, None)
        return

    # 2. توليد وتجهيز الملف وإرساله للطالب بنجاح
    try:
        if out_type == "docx":
            file_obj = create_pro_academic_docx(solution, custom_title=doc_title, filename=safe_filename)
            await callback.message.answer_document(
                file_obj,
                caption=f"📄 **تم إعداد ملف Word باسم:** `{safe_filename}.docx`",
                parse_mode="Markdown"
            )
        elif out_type == "pdf":
            file_obj = create_pro_academic_pdf(solution, custom_title=doc_title, filename=safe_filename)
            await callback.message.answer_document(
                file_obj,
                caption=f"📑 **تم إعداد ملف PDF باسم:** `{safe_filename}.pdf`",
                parse_mode="Markdown"
            )
        else:
            if len(solution) > 4000:
                for chunk in [solution[i:i+4000] for i in range(0, len(solution), 4000)]:
                    await callback.message.answer(chunk)
            else:
                await callback.message.answer(solution)

        # 3. الخصم فقط بعد تسليم الملف بنجاح
        await db.deduct_credit(user_id)
        credits = await db.get_or_create_user(user_id)
        await callback.message.answer(f"✅ الرصيد المتبقي في حسابك: **{credits}** نقاط.")
        await wait_msg.delete()

    except Exception as e:
        print(f"Delivery Error: {e}")
        await wait_msg.edit_text(
            "⚠️ **حدث خطأ أثناء تصدير المستند.**\n"
            "💡 رصيدك محفوظ بالكامل ولم يُخصم منه شيء. يرجى المحاولة مرة أخرى أو اختيار صيغة مختلفة."
        )

    USER_PENDING_TASKS.pop(user_id, None)


async def main():
    await db.init_db()
    print("🤖 البوت يعمل بنظام تنسيق المستندات الأكاديمية ودعم النصوص العربية المتقدم...")
    await dp.start_polling(bot)


if __name__ == "__main__":
    asyncio.run(main())